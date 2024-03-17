"""
App to save screenshots and format a document.
"""
import sys
import os
import datetime

import pyautogui
from PyQt5 import QtWidgets
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QWidget
from PyQt5.QtCore import QObject, Qt
from pynput import mouse
from docx import Document
from docx.shared import Inches

from app import Ui_Form
from captureWindow import select_window

HOME_DIRECTORY = os.path.join(
    os.path.expanduser("~\\Documents"), "StepCapture")

DISABLE_BUTTON_STYLE = '''
        QPushButton {
            color: #808080;  /* Grey text color */
            background-color: #f0f0f0;  /* Light grey background */
            border: 1px solid #d3d3d3;  /* Light grey border */
            border-radius: 5px;  /* Rounded corners */
            padding: 5px 10px;  /* Padding */
        }
        QPushButton:disabled {
            color: #b0b0b0;  /* Darker grey text color for disabled state */
            background-color: #e0e0e0;  /* Lighter grey background for disabled state */
            border: 1px solid #d3d3d3;  /* Light grey border for disabled state */
        }
    '''


def ensure_directory(directory):
    """
    Ensure the directory exists, creating it if necessary.

    Args:
        directory (str): The directory path.
    """
    if not os.path.exists(directory):
        os.makedirs(directory)


class MainWindow(QWidget, QObject):
    def __init__(self):
        super().__init__()
        self.ui = Ui_Form()
        self.ui.setupUi(self)
        self.setup_ui_events()
        self.setMouseTracking(True)
        self.mouseListener = None
        self.home_directory = None
        self.capture_area = None
        self.take_screenshot = False
        self.doc = None
        ensure_directory(HOME_DIRECTORY)

        # Set the window flag to keep the window on top
        self.setWindowFlags(self.windowFlags() | Qt.WindowStaysOnTopHint)
        self.setWindowIcon(QIcon('D:\dev\dev.py\step-capture\stepCapture.ico'))

    def setup_ui_events(self):
        """
        Set up UI events.
        """
        self.ui.selectAreaPushButton.clicked.connect(self.select_area)
        self.ui.startCapturePushButton.clicked.connect(self.start_capture)
        self.ui.createDocumentPushButton.clicked.connect(self.create_document)

    def select_area(self):
        """
        Select the area on the screen for capture.
        """
        # self.ui.selectAreaPushButton.setStyleSheet(DISABLE_BUTTON_STYLE)
        self.ui.logLineEdit.setText(
            "Click TopRight > BottomLeft corner of Capture window.")
        self.ui.logLineEdit.repaint()
        print("Select Area button clicked")
        self.capture_area = select_window()
        print(f"Area: {self.capture_area}")

        self.ui.logLineEdit.setText(
            "Click click Start Capture, or Select Area to select again.")

    def start_capture(self):
        """
        Start capturing screenshots.
        """
        print("Start Capture button clicked")
        try:
            self.mouseListener = mouse.Listener(on_click=self.on_click)
            self.mouseListener.start()
            self.create_document()
            print("Take capture from window")
            self.take_screenshot = True
            self.ui.startCapturePushButton.setEnabled(False)
            # self.ui.startCapturePushButton.setStyleSheet(DISABLE_BUTTON_STYLE)
            self.ui.logLineEdit.setText(
                "Click inside selected area to capture.")

        except Exception as ex:
            print(f"Start Capture Error: {ex}")

    def on_click(self, x, y, button, pressed):
        """
        Handle mouse clicks to capture screenshots.
        """
        if self.take_screenshot and pressed and self.is_within_capture_area(x, y):
            print(
                f"Save capture state - {self.capture_area} - Ctrl+Clicks pressed")
            self.screenshot()

    def screenshot(self):
        """
        Capture a screenshot of the selected area.
        """
        # Get the screen geometry
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = os.path.join(self.home_directory,
                                f"screenshot_{timestamp}.png")
        (x1, y1), (x2, y2) = self.capture_area
        width = x2 - x1
        height = y2 - y1
        screenshot = pyautogui.screenshot(region=(x1, y1, width, height))
        self.doc.add_paragraph("")
        self.doc.add_paragraph(os.path.basename(filename))
        screenshot.save(filename)
        self.doc.add_picture(filename, width=Inches(5))

    def is_within_capture_area(self, x, y):
        """
        Check if the mouse click is within the capture area.
        """
        (x1, y1), (x2, y2) = self.capture_area
        return x1 <= x <= x2 and y1 <= y <= y2

    def create_document(self):
        """
        Create a Word document with captured screenshots.
        """
        if self.doc is None:
            print("Create Document button clicked")

            # Create a folder to store screenshots with a timestamp as its name
            timestamp = f'step-Capture-{datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")}'
            self.home_directory = os.path.join(HOME_DIRECTORY, timestamp)
            ensure_directory(self.home_directory)

            self.doc = Document()
            self.doc.add_heading('Main Heading [EDIT]', level=1)
            self.doc.add_paragraph("Follow below steps to....")
            self.doc.add_paragraph("...")
            self.ui.createDocumentPushButton.setText("Save Document")
            self.ui.logLineEdit.setText(
                "Click Save Document to save and start new.")
        else:
            print("Save Document button clicked")
            self.take_screenshot = False
            self.mouseListener.stop()
            self.doc.save(
                f'{os.path.join(self.home_directory, os.path.basename(self.home_directory))}.docx')
            self.doc = None
            self.ui.startCapturePushButton.setEnabled(True)
            # self.ui.startCapturePushButton.setStyleSheet(DISABLE_BUTTON_STYLE)
            self.ui.createDocumentPushButton.setText("Create Document")
            self.ui.logLineEdit.setText(
                "Select new area or start capture to start new.")

    def closeEvent(self, event):
        # Stop the mouse listener when the window is closed
        print("Closing the window.")
        # Save the document
        if self.doc is not None:
            self.doc.save(
                f'{os.path.join(self.home_directory, os.path.basename(self.home_directory))}.docx')
        event.accept()
        if self.home_directory:
            os.startfile(self.home_directory)


if __name__ == "__main__":
    try:
        app = QtWidgets.QApplication(sys.argv)
        mainWindow = MainWindow()
        mainWindow.show()
        app.installEventFilter(mainWindow)
        sys.exit(app.exec_())
    except Exception as ex:
        print(f"MAIN Error: {ex}")
